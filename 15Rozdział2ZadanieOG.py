import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_invitations(guests_file, output_file):
    # Load the guest names from the file
    with open(guests_file, 'r') as file:
        guests = file.readlines()

    # Create a new Word document
    doc = docx.Document()

    for guest in guests:
        guest = guest.strip()  # Remove any leading/trailing whitespace

        # Add invitation text
        doc.add_paragraph('It would be a pleasure to have the company of', style='Normal')
        para = doc.add_paragraph(guest, style='Normal')
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the guest name
        doc.add_paragraph('at 11010 Memory Lane on the Evening of April 1st', style='Normal')
        doc.add_paragraph('at 7 o’clock', style='Normal')

        # Add a page break after each invitation, except for the last one
        if guest != guests[-1].strip():
            doc.add_page_break()

    # Save the document
    doc.save(output_file)

create_invitations('guests.txt', 'invitations.docx')
