# generate-invitations-docx.py
# Generates formal invitations for a list of guests and saves them to a Word document,
# placing each invitation on a separate page.

import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_invitations(guests_file, output_file):
    """
    Creates a Word document with formal invitations for each guest listed in a text file.

    Args:
    guests_file (str): The path to the text file containing guest names, one per line.
    output_file (str): The path for the output Word document containing the invitations.
    """
    # Load guest names from the provided text file.
    with open(guests_file, 'r') as file:
        guests = file.readlines()

    # Initialize a new Word document.
    doc = docx.Document()

    # Generate an invitation for each guest.
    for guest in guests:
        guest = guest.strip()  # Clean guest name from trailing whitespaces or newlines.

        # Add the invitation text, with the guest's name centered.
        doc.add_paragraph('It would be a pleasure to have the company of', style='Normal')
        para = doc.add_paragraph(guest, style='Normal')
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the guest name.
        doc.add_paragraph('at 11010 Memory Lane on the Evening of April 1st', style='Normal')
        doc.add_paragraph('at 7 o’clock', style='Normal')

        # Insert a page break after each invitation, except for the last one.
        if guest != guests[-1].strip():
            doc.add_page_break()

    # Save the generated invitations to the specified Word document.
    doc.save(output_file)


# Example function call to create invitations.
create_invitations('guests.txt', 'invitations.docx')
